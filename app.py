#!/usr/bin/env python3
"""
Asset Return Form - Flask Web Application
Run: python app.py
Then open: http://localhost:5000
"""

from flask import Flask, render_template, request, send_file, jsonify, abort
from docx import Document
import pandas as pd
from threading import Lock
import os
import shutil
from datetime import datetime
import subprocess
import sys
import shlex
from werkzeug.utils import secure_filename

# Optional converter import (docx2pdf on Windows)
try:
    from docx2pdf import convert as docx2pdf_convert  # type: ignore
except Exception:
    docx2pdf_convert = None

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "Asset_Return_Form_Template.docx")
OUTPUT_DIR = os.path.join(BASE_DIR, "Asset_Return_Forms")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Field definitions: key -> (label, section, type)
FIELDS = {
    # Section 1
    "employee_name":        ("Employee Name", "1. Employee Details", "text"),
    "employee_id":          ("Employee ID", "1. Employee Details", "text"),
    "department":           ("Department", "1. Employee Details", "text"),
    "designation":          ("Designation", "1. Employee Details", "text"),
    "email_id":             ("Email ID", "1. Employee Details", "text"),
    "contact_number":       ("Contact Number", "1. Employee Details", "text"),
    "reporting_manager":    ("Reporting Manager", "1. Employee Details", "text"),
    "last_working_day":     ("Last Working Day", "1. Employee Details", "text"),

    # Section 2
    "asset_name":           ("Asset Name", "2. Asset Details", "text"),
    "asset_id":             ("Asset ID / Tag Number", "2. Asset Details", "text"),
    "serial_number":        ("Serial Number", "2. Asset Details", "text"),
    "asset_type":           ("Asset Type", "2. Asset Details", "text"),
    "issued_date":          ("Issued Date", "2. Asset Details", "text"),
    "condition_at_return":  ("Condition at Return", "2. Asset Details", "text"),
    "asset_remarks":        ("Remarks", "2. Asset Details", "text"),

    # Section 3
    "charger_qty":          ("Charger / Adapter - Quantity", "3. Accessories", "text"),
    "charger_condition":    ("Charger / Adapter - Condition", "3. Accessories", "text"),
    "bag_qty":              ("Bag - Quantity", "3. Accessories", "text"),
    "bag_condition":        ("Bag - Condition", "3. Accessories", "text"),

    # Section 4
    "data_backed_up":       ("Official data backed up", "4. Verification", "select"),
    "email_disabled":       ("Email access disabled", "4. Verification", "select"),
    "vpn_revoked":          ("VPN / System access revoked", "4. Verification", "select"),
    "asset_register":       ("Asset updated in register", "4. Verification", "select"),

    # Section 5
    "decl_employee_name":   ("Employee Name (Declaration)", "5. Declaration", "text"),
    "decl_date":            ("Declaration Date", "5. Declaration", "text"),

    # Section 6
    "it_received_by":       ("Asset Received By", "6. IT Verification", "text"),
    "it_designation":       ("Designation", "6. IT Verification", "text"),
    "it_date":              ("Date", "6. IT Verification", "text"),
}

# -------------------------
# Employee Excel loader (hardcoded employee_id column)
# -------------------------
EMPLOYEE_XLSX = os.path.join(BASE_DIR, "data", "employee_info.xlsx")

_employee_cache = None
_employee_cache_lock = Lock()
_employee_cache_mtime = None

def load_employee_data(force_reload: bool = False):
    """
    Load employee Excel into a dict keyed by employee_id.
    Returns: dict { employee_id_str: {form_key: value, ...}, ... }
    - Uses file mtime to avoid reloading on every request.
    - If Excel columns differ from form keys, provide mapping in excel_to_form.
    """
    global _employee_cache, _employee_cache_mtime

    try:
        mtime = os.path.getmtime(EMPLOYEE_XLSX)
    except FileNotFoundError:
        return {}

    with _employee_cache_lock:
        if _employee_cache is not None and not force_reload and _employee_cache_mtime == mtime:
            return _employee_cache

        # Read Excel into DataFrame (all values as strings)
        df = pd.read_excel(EMPLOYEE_XLSX, engine="openpyxl", dtype=str).fillna("")

        # Normalize column names (trim whitespace)
        df.columns = [c.strip() for c in df.columns]

        # If your sheet header is "Employee ID" (with space), map it to internal key "employee_id"
        # Change the left-hand string if your header is different.
        if "Employee ID" in df.columns and "employee_id" not in df.columns:
            df = df.rename(columns={"Employee ID": "employee_id"})

        # Optional: map other human-friendly headers to internal form keys
        excel_to_form = {
            "Employee Name": "employee_name",
            "Employee ID": "employee_id",
            "Department": "department",
            "Designation": "designation",
            "Email ID": "email_id",
            "Contact Number": "contact_number",
            "Reporting Manager": "reporting_manager",
            "Last Working Day": "last_working_day",

            "Asset Name": "asset_name",
            "Asset ID / Tag Number": "asset_id",
            "Serial Number": "serial_number",
            "Asset Type": "asset_type",
            "Issued Date": "issued_date",
            "Condition at Return": "condition_at_return",
            "Remarks": "asset_remarks",

            "Charger / Adapter - Quantity": "charger_qty",
            "Charger / Adapter - Condition": "charger_condition",
            "Bag - Quantity": "bag_qty",
            "Bag - Condition": "bag_condition",

            "Official data backed up": "data_backed_up",
            "Email access disabled": "email_disabled",
            "VPN / System access revoked": "vpn_revoked",
            "Asset updated in register": "asset_register",

            "Employee Name (Declaration)": "decl_employee_name",
            "Declaration Date": "decl_date",

            "Asset Received By": "it_received_by",
            "Designation (IT)": "it_designation",
            "Date (IT)": "it_date"
        }

        
        if excel_to_form:
            df = df.rename(columns=excel_to_form)

        cache = {}
        # Iterate rows and build a record for each employee
        for _, row in df.iterrows():
            emp_id = str(row.get("employee_id", "")).strip()
            if not emp_id:
                continue

            record = {}
            for key in FIELDS.keys():
                # Prefer a column that exactly matches the internal key
                if key in df.columns:
                    record[key] = str(row.get(key, "")).strip()
                else:
                    # Fallback: use the human label from FIELDS (FIELDS[key][0])
                    label = FIELDS[key][0]
                    if label in df.columns:
                        record[key] = str(row.get(label, "")).strip()
                    else:
                        record[key] = ""
            cache[emp_id] = record

        _employee_cache = cache
        _employee_cache_mtime = mtime
        return _employee_cache


def get_sections():
    """Group fields by section for template rendering."""
    sections = {}
    for key, (label, section, ftype) in FIELDS.items():
        sections.setdefault(section, []).append((key, label, ftype))
    return sections


def fill_template(data, output_path):
    """Fill the Word template with provided data."""
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found at {TEMPLATE_PATH}")

    shutil.copy(TEMPLATE_PATH, output_path)
    doc = Document(output_path)

    label_to_key = {
        "User Name": "employee_name",
        "Employee ID": "employee_id",
        "Department": "department",
        "Designation": "designation",
        "Official Email ID": "email_id",
        "Number": "contact_number",
        "Manager Name": "reporting_manager",
        "Last Working Day": "last_working_day",

        "Asset Name": "asset_name",
        "Asset ID": "asset_id",
        "Serial Number": "serial_number",
        "Asset Type": "asset_type",
        "Issued Date": "issued_date",
        "Condition at Return": "condition_at_return",
        "Remarks": "asset_remarks",

        "Charger / Adapter": ("charger_qty", "charger_condition"),
        "Bag": ("bag_qty", "bag_condition"),

        "Official data backed up": "data_backed_up",
        "Email access disabled": "email_disabled",
        "VPN / System access revoked": "vpn_revoked",
        "Asset updated in asset register": "asset_register",

        "Employee Name:": "decl_employee_name",
        "Asset Received By:": "it_received_by",
        "Designation:": "it_designation",
        "Date:": "it_date",
    }

    # Handle top-level Date in paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "Date:" or "**Date:**" in text:
            if data.get("decl_date"):
                para.clear()
                para.add_run(f"Date: {data['decl_date']}")
            break

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue

            first_text = cells[0].text.strip()

            for label, mapping in label_to_key.items():
                if label.lower() in first_text.lower():
                    if isinstance(mapping, tuple):
                        qty_key, cond_key = mapping
                        if len(cells) >= 2 and data.get(qty_key):
                            cells[1].text = data[qty_key]
                        if len(cells) >= 3 and data.get(cond_key):
                            cells[2].text = data[cond_key]
                    else:
                        if len(cells) >= 2 and data.get(mapping):
                            cells[1].text = data[mapping]
                    break

    doc.save(output_path)
    return output_path


def convert_docx_to_pdf(docx_path, out_dir):
    """
    Convert a .docx to .pdf and return the pdf path.
    Tries platform-appropriate converters:
      - Windows: docx2pdf (if installed) or win32com fallback
      - Linux/macOS: LibreOffice 'soffice' headless
    """
    base = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(out_dir, f"{base}.pdf")

    # Remove existing PDF to avoid stale files
    if os.path.exists(pdf_path):
        try:
            os.remove(pdf_path)
        except Exception:
            pass

    if sys.platform == "win32":
        # Prefer docx2pdf if available
        if docx2pdf_convert:
            # docx2pdf can accept (input, output) or a folder; use file->file
            docx2pdf_convert(docx_path, pdf_path)
        else:
            # Try win32com automation (requires pywin32)
            try:
                import win32com.client  # type: ignore
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(docx_path))
                # 17 = wdFormatPDF
                doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
                doc.Close(False)
                word.Quit()
            except Exception as e:
                raise RuntimeError("No converter available on Windows: " + str(e))
    else:
        # Use LibreOffice headless conversion
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            raise RuntimeError("LibreOffice (soffice) not found. Install it or provide a converter.")
        cmd = f'{shlex.quote(soffice)} --headless --convert-to pdf --outdir {shlex.quote(out_dir)} {shlex.quote(docx_path)}'
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.strip()}")
        # Verify output exists
        if not os.path.exists(pdf_path):
            # Try to find created pdf in out_dir
            created = None
            for f in os.listdir(out_dir):
                if f.lower().startswith(base.lower()) and f.lower().endswith(".pdf"):
                    created = os.path.join(out_dir, f)
                    break
            if created:
                pdf_path = created
            else:
                raise RuntimeError("PDF not found after conversion.")
    return pdf_path


def safe_path_in_output(filename):
    """Return absolute path for filename inside OUTPUT_DIR and validate it."""
    filename = secure_filename(filename)
    path = os.path.abspath(os.path.join(OUTPUT_DIR, filename))
    if not path.startswith(os.path.abspath(OUTPUT_DIR) + os.sep) and os.path.abspath(OUTPUT_DIR) != path:
        # Prevent path traversal
        raise ValueError("Invalid filename")
    return path


@app.route("/")
def index():
    return render_template("index.html", sections=get_sections())


@app.route("/generate", methods=["POST"])
def generate():
    data = {k: request.form.get(k, "").strip() for k in FIELDS.keys()}

    emp_name = data.get("employee_name", "").strip() or "Unknown_Employee"
    emp_name = "".join(c for c in emp_name if c.isalnum() or c in " -_").rstrip().replace(" ", "_")

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"Asset_Return_{emp_name}_{ts}.docx"
    docx_path = os.path.join(OUTPUT_DIR, docx_filename)

    try:
        fill_template(data, docx_path)
    except Exception as e:
        return jsonify({"success": False, "message": f"Template fill failed: {e}"}), 500

    pdf_filename = None
    try:
        pdf_path = convert_docx_to_pdf(docx_path, OUTPUT_DIR)
        pdf_filename = os.path.basename(pdf_path)
    except Exception as e:
        # Conversion failed; log and continue returning docx as fallback
        print("PDF conversion failed:", e)

    # Build response URLs (prefer PDF for preview/download if available)
    if pdf_filename:
        download_url = f"/download/{pdf_filename}"
        preview_url = f"/preview/{pdf_filename}"
    else:
        download_url = f"/download/{docx_filename}"
        preview_url = f"/preview/{docx_filename}"

    return jsonify({
        "success": True,
        "docx_filename": docx_filename,
        "pdf_filename": pdf_filename,
        "download_url": download_url,
        "preview_url": preview_url,
        "print_url": f"/print/{docx_filename}"
    })


@app.route("/download/<path:filename>")
def download(filename):
    try:
        path = safe_path_in_output(filename)
    except ValueError:
        abort(400, "Invalid filename")

    if os.path.exists(path):
        mimetype = "application/pdf" if filename.lower().endswith(".pdf") else None
        return send_file(path, as_attachment=True, mimetype=mimetype)
    return "File not found", 404


@app.route("/preview/<path:filename>")
def preview(filename):
    try:
        path = safe_path_in_output(filename)
    except ValueError:
        abort(400, "Invalid filename")

    if os.path.exists(path):
        mimetype = "application/pdf" if filename.lower().endswith(".pdf") else None
        # as_attachment=False so browser will try to render inline
        return send_file(path, as_attachment=False, mimetype=mimetype)
    return "File not found", 404


@app.route("/lookup", methods=["GET"])
def lookup_emp():
    """
    GET /lookup?employee_id=12345
    Returns JSON { success: True, data: {form_key: value, ...} }
    """
    emp_id = request.args.get("employee_id", "").strip()
    if not emp_id:
        return jsonify({"success": False, "message": "employee_id required"}), 400

    data = load_employee_data()
    record = data.get(emp_id)

    # case-insensitive fallback
    if not record:
        for k, v in data.items():
            if k.lower() == emp_id.lower():
                record = v
                break

    if not record:
        return jsonify({"success": False, "message": "Employee not found"}), 404

    return jsonify({"success": True, "data": record})


@app.route("/print/<path:filename>")
def print_file(filename):
    """
    Server-side print fallback. This attempts to send the DOCX to a server-side printer.
    Recommended flow is client-side printing of the PDF (browser print dialog).
    """
    try:
        path = safe_path_in_output(filename)
    except ValueError:
        abort(400, "Invalid filename")

    if not os.path.exists(path):
        return jsonify({"success": False, "message": "File not found"}), 404

    try:
        # If a PDF exists with same base name, prefer printing the PDF
        base = os.path.splitext(os.path.basename(path))[0]
        pdf_candidate = os.path.join(OUTPUT_DIR, f"{base}.pdf")
        target_path = pdf_candidate if os.path.exists(pdf_candidate) else path

        if sys.platform == "win32":
            # Try Word automation first for DOCX; for PDF use os.startfile
            if target_path.lower().endswith(".pdf"):
                os.startfile(target_path, "print")
            else:
                try:
                    import win32com.client  # type: ignore
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(target_path))
                    doc.PrintOut()
                    doc.Close(SaveChanges=False)
                    word.Quit()
                except Exception:
                    os.startfile(target_path, "print")
        elif sys.platform == "darwin":
            subprocess.run(["lp", target_path], check=False)
        else:
            # Linux/Unix: use lp or lpr
            result = subprocess.run(["lp", target_path], capture_output=True, text=True)
            if result.returncode != 0:
                subprocess.run(["lpr", target_path], check=False)

        return jsonify({"success": True, "message": "Sent to printer!"})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500


if __name__ == "__main__":
    print("=" * 55)
    print("  🚀 Asset Return Form - Web App")
    print("  Open your browser: http://localhost:5000")
    print("  Press Ctrl+C to stop")
    print("=" * 55)
    app.run(host="0.0.0.0", port=5000, debug=True)
